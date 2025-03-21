from docx import Document

def export_to_word(title, spec_text, improvement_text, output_file="produkt_spec.docx"):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_heading("Teknisk Specifikation", level=2)
    doc.add_paragraph(spec_text)
    doc.add_heading("Föreslagna Förbättringar", level=2)
    doc.add_paragraph(improvement_text)
    doc.save(output_file)